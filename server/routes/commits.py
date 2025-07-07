from fastapi import APIRouter, Form, Depends, HTTPException
from sqlalchemy.orm import Session
from datetime import datetime
import os
import uuid
import pathlib
from server.models import Log, FileVersion, MergeHistory, User ,Commit, Snapshot, AccessControl , FileVersionHistory
from server.db import get_db
from server.dependencies import get_current_user
from server.schemas import UserOut
from server.routes.repositories import compute_file_hash
from fpdf import FPDF
from server.utils.diff_utils import generate_diff
from docx import Document

router = APIRouter(prefix="/repositories", tags=["Commits"])
import re

def strip_version_suffix(filename):
    # e.g., a_v1.docs → a.docs
    base, ext = os.path.splitext(filename)
    base = re.sub(r'_v\d+$', '', base)
    return base + ext
@router.post("/{repo_id}/commit")
def create_commit(
    repo_id: int,
    message: str = Form(...),
    filename: str = Form(...),
    content: str = Form(...),
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    #   1. Access check
    access = db.query(AccessControl).filter_by(
        user_id=current_user.id,
        repository_id=repo_id
    ).first()
    if not access or access.role not in ("admin", "editor","collaborator"):
        raise HTTPException(status_code=403, detail="Permission denied")

    #   2. Normalize path
    safe_rel_path = pathlib.Path(filename).as_posix().lstrip("/\\")
    if not safe_rel_path.endswith(".txt"):
        safe_rel_path = os.path.splitext(safe_rel_path)[0] + ".txt"

    full_path = os.path.join(f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}", safe_rel_path)
    ext = ".txt"


    #   3. If it's a text file, generate diff
    if ext == ".txt":
        original = ""
        if os.path.exists(full_path):
            with open(full_path, "r", encoding="utf-8") as f:
                original = f.read()

        diff = generate_diff(original, content)

        commit = Commit(
            repo_id=repo_id,
            author_id=current_user.id,
            message=message,
            original_filename=safe_rel_path,
            versioned_filename=None,
            snapshot_path=None,
            diff_text=diff,
            status="proposed"
        )
        db.add(commit)
        db.commit()
        db.refresh(commit)

        return {
            "commit_id": commit.id,
            "message": f"Text commit created with diff for {safe_rel_path}",
            "type": "diff"
        }
@router.post("/{repo_id}/revert/{commit_id}")
def revert_commit(
    repo_id: int,
    commit_id: int,
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    from server.utils.diff_utils import apply_diff

    # ───── 1. Access Check ─────
    access = db.query(AccessControl).filter_by(
        user_id=current_user.id, repository_id=repo_id
    ).first()
    if not access or access.role not in ("admin", "editor", "collaborator"):
        raise HTTPException(status_code=403, detail="Permission denied")

    # ───── 2. Fetch Commit ─────
    commit = db.query(Commit).filter_by(id=commit_id, repo_id=repo_id).first()
    if not commit:
        raise HTTPException(status_code=404, detail="Commit not found")

    repo_root = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}"
    version_dir = os.path.join(repo_root, "versions")

    # ───── 3. Handle Revert of Merged Commit (From Two Commits Merge) ─────
    merge_entry = db.query(MergeHistory).filter_by(result_commit_id=commit.id).first()
    if merge_entry and commit.status == "merged":
        version_file = commit.versioned_filename
        version_path = os.path.join(version_dir, version_file)

        if version_file and os.path.exists(version_path):
            os.remove(version_path)

        commit.status = "proposed"
        commit.versioned_filename = None

        # Remove file version history
        db.query(FileVersionHistory).filter_by(commit_id=commit.id).delete()
        db.commit()

        db.add(Log(
            user_id=current_user.id,
            repo_id=repo_id,
            commit_id=commit.id,
            action="revert_merge",
            description=f"Reverted merge of two commits (commit #{commit.id})",
            timestamp=datetime.utcnow()
        ))
        db.commit()

        return {"message": f"Reverted merged version file from commit #{commit.id}"}

    # ───── 4. Handle Revert of Commit Merged with Original File ─────
    fvh = db.query(FileVersionHistory).filter_by(commit_id=commit.id).first()
    if fvh and commit.status == "merged":
        version_file = fvh.versioned_filename
        version_path = os.path.join(version_dir, version_file)

        if version_file and os.path.exists(version_path):
            os.remove(version_path)

        commit.status = "proposed"
        commit.versioned_filename = None

        db.query(FileVersionHistory).filter_by(commit_id=commit.id).delete()
        db.commit()

        db.add(Log(
            user_id=current_user.id,
            repo_id=repo_id,
            commit_id=commit.id,
            action="revert_merge",
            description=f"Reverted merged commit with original file (commit #{commit.id})",
            timestamp=datetime.utcnow()
        ))
        db.commit()

        return {"message": f"Reverted merged commit with original file from commit #{commit.id}"}

    # ───── 5. Handle Revert of Proposed Commit (diff restore) ─────
    if commit.status == "proposed":
        original_path = os.path.join(repo_root, commit.original_filename)

        if os.path.exists(original_path):
            raise HTTPException(status_code=400, detail="File already exists. Cannot restore.")

        if not commit.diff_text:
            raise HTTPException(status_code=400, detail="No diff available to restore")

        try:
            restored_content = apply_diff("", commit.diff_text)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to apply diff: {str(e)}")

        os.makedirs(os.path.dirname(original_path), exist_ok=True)
        with open(original_path, "w", encoding="utf-8") as f:
            f.write(restored_content)

        db.add(Log(
            user_id=current_user.id,
            repo_id=repo_id,
            commit_id=commit.id,
            action="revert",
            description=f"Restored file from proposed commit #{commit.id}",
            timestamp=datetime.utcnow()
        ))
        db.commit()

        return {"message": f"Proposed commit #{commit.id} reverted. File restored."}

    raise HTTPException(status_code=400, detail="Unsupported commit status for revert")
@router.post("/{repo_id}/revert_version/{version_filename}")
def revert_merged_version_file(
    repo_id: int,
    version_filename: str,
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    # Ensure only admin can revert
    access = db.query(AccessControl).filter_by(user_id=current_user.id, repository_id=repo_id).first()
    if not access or access.role != "admin":
        raise HTTPException(status_code=403, detail="Permission denied")

    # Find commit with this versioned filename
    commit = db.query(Commit).filter_by(
        repo_id=repo_id,
        versioned_filename=version_filename,
        status="merged"
    ).first()

    if not commit:
        raise HTTPException(status_code=404, detail="No merged commit found for this version file.")

    # Check if it was created via version-version merge (2 parents)
    merge_history = db.query(MergeHistory).filter_by(result_commit_id=commit.id).first()
    if not merge_history:
        raise HTTPException(status_code=400, detail="This file was not created by merging two version files.")

    version_path = os.path.join(
        f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}/versions",
        version_filename
    )

    # Delete the version file
    if os.path.exists(version_path):
        os.remove(version_path)

    # Reset status if needed
    commit.status = "proposed"
    commit.versioned_filename = None
    db.commit()

    # Delete merge history entry
    db.query(MergeHistory).filter_by(result_commit_id=commit.id).delete()
    db.commit()

    # Delete version history
    db.query(FileVersionHistory).filter_by(commit_id=commit.id).delete()
    db.commit()

    # Log
    db.add(Log(
        user_id=current_user.id,
        repo_id=repo_id,
        commit_id=commit.id,
        action="revert_version_merge",
        description=f"Reverted version-version merge file {version_filename}",
        timestamp=datetime.utcnow()
    ))
    db.commit()

    return {"message": f"Reverted version file {version_filename} created from merging two versions."}
@router.post("/{repo_id}/revert_version")
def revert_merged_version_file(
    repo_id: int,
    versioned_filename: str = Form(...),
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    import os
    safe_versioned_filename = os.path.basename(versioned_filename)

    # Step 1: Check version history
    fvh = db.query(FileVersionHistory).filter_by(
        repo_id=repo_id,
        versioned_filename=safe_versioned_filename
    ).first()

    if not fvh:
        raise HTTPException(404, "Version file not found in version history")

    if not fvh.commit_id:
        raise HTTPException(400, "This version is not linked to any commit")

    # Step 2: Check commit status
    commit = db.query(Commit).filter_by(id=fvh.commit_id).first()
    if not commit or commit.status != "merged":
        raise HTTPException(400, "This file was not created via merge and cannot be reverted.")

    # Step 3: Delete the actual file
    version_path = os.path.join(f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}/versions", safe_versioned_filename)
    if os.path.exists(version_path):
        os.remove(version_path)

    # Step 4: Reset commit
    commit.status = "proposed"
    commit.versioned_filename = None
    db.commit()

    # Step 5: Clean up version history
    db.query(FileVersionHistory).filter_by(commit_id=commit.id).delete()
    db.commit()

    # Step 6: Log
    db.add(Log(
        user_id=current_user.id,
        repo_id=repo_id,
        commit_id=commit.id,
        action="revert",
        description=f"Reverted versioned file {safe_versioned_filename}",
        timestamp=datetime.utcnow()
    ))
    db.commit()

    return {"message": f"Reverted version file {safe_versioned_filename}"}


from fastapi.responses import FileResponse

def get_next_flat_version_filename(folder: str, filename: str) -> str:
    """
    Extracts base from given filename and finds next flat version filename.
    Handles both commit merges and version-version merges.
    """
    name, ext = os.path.splitext(os.path.basename(filename))
    base = re.sub(r"_v\d+$", "", name)

    pattern = re.compile(rf"^{re.escape(base)}_v(\d+){re.escape(ext)}$")
    max_version = 0

    for f in os.listdir(folder):
        match = pattern.match(f)
        if match:
            max_version = max(max_version, int(match.group(1)))

    return f"{base}_v{max_version + 1}{ext}"


from fastapi import Body
from typing import List, Optional


@router.post("/{repo_id}/merge")
def merge_multiple_commits_or_versions(
    repo_id: int,
    file_name: Optional[str] = Body(None),
    commit_ids: List[int] = Body(...),
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    #   Access check
    access = db.query(AccessControl).filter_by(
        user_id=current_user.id,
        repository_id=repo_id
    ).first()
    if not access or access.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can merge.")

    if len(commit_ids) < 2:
        raise HTTPException(status_code=400, detail="At least 2 commits required to merge.")

    commits = db.query(Commit).filter(Commit.id.in_(commit_ids)).all()

    if len(commits) != len(commit_ids):
        raise HTTPException(status_code=404, detail="Some commits not found.")

    #   All commits must be from the same file lineage
    filenames = {strip_version_suffix(c.original_filename or c.versioned_filename) for c in commits}
    if len(filenames) != 1:
        raise HTTPException(status_code=400, detail="Commits must belong to the same original file.")

    #   Combine content from all commits (you can choose smarter merge logic)
    merged_content = ""
    from server.utils.diff_utils import apply_diff

    repo_root = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}"

    for commit in commits:
        content = ""
        if commit.diff_text:
            original_path = os.path.join(repo_root, commit.original_filename)
            base_content = ""
            if os.path.exists(original_path):
                with open(original_path, "r", encoding="utf-8") as f:
                    base_content = f.read()
            content = apply_diff(base_content, commit.diff_text)

        elif commit.snapshot_path and os.path.exists(commit.snapshot_path):
            with open(commit.snapshot_path, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            raise HTTPException(status_code=404, detail=f"No valid snapshot or diff for commit {commit.id}")

        merged_content += f"\n\n# Commit {commit.id}\n{content.strip()}\n"


    #   Save merged content to new version file
    repo_folder = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}/versions"
    os.makedirs(repo_folder, exist_ok=True)

    base_filename = commits[0].original_filename or commits[0].versioned_filename
    versioned_filename = get_next_flat_version_filename(repo_folder, base_filename)
    version_path = os.path.join(repo_folder, versioned_filename)

    with open(version_path, "w", encoding="utf-8") as out:
        out.write(merged_content)

    #   Register merged commit
    merged_commit = Commit(
        repo_id=repo_id,
        author_id=current_user.id,
        message=f"Merged commits {', '.join(map(str, commit_ids))}",
        original_filename=base_filename,
        versioned_filename=versioned_filename,
        snapshot_path=version_path,
        status="merged"
    )
    db.add(merged_commit)
    db.commit()
    db.refresh(merged_commit)
    # ✅ Insert into MergeHistory if exactly 2 commits merged
    if len(commit_ids) == 2:
        base_commit_id = commit_ids[0]
        merged_commit_id = commit_ids[1]

        merge_entry = MergeHistory(
            repo_id=repo_id,
            base_commit_id=base_commit_id,
            merged_commit_id=merged_commit_id,
            result_commit_id=merged_commit.id,
            result_filename=versioned_filename,
            merged_by=current_user.id,
            timestamp=datetime.utcnow()
        )
        db.add(merge_entry)
        db.commit()

    #   Update all merged commits’ status
    for c in commits:
        c.status = "merged"
    db.commit()

    #   Log
    log = Log(
        user_id=current_user.id,
        repo_id=repo_id,
        commit_id=merged_commit.id,
        action="merge",
        description=f"Merged commits {commit_ids} into {versioned_filename}",
        timestamp=datetime.utcnow()
    )
    db.add(log)
    db.commit()
    merged_content = merged_content.strip()
    # 1. Determine or create FileVersion row
    fv = db.query(FileVersion).filter_by(repo_id=repo_id, original_filename=base_filename).first()
    if not fv:
        fv = FileVersion(repo_id=repo_id, original_filename=base_filename, latest_version=1)
        db.add(fv)
        db.commit()
        db.refresh(fv)
    else:
        fv.latest_version += 1
        db.commit()

    # 2. Add to FileVersionHistory
    fvh = FileVersionHistory(
        repo_id=repo_id,
        original_filename=base_filename,
        version_no=fv.latest_version,
        commit_id=commit.id,
        versioned_filename=versioned_filename,
        snapshot_path=version_path,
        merged_by=current_user.id,
    )
    db.add(fvh)
    db.commit()

    return {
        "message": f"Merged into {versioned_filename}",
        "versioned_filename": versioned_filename,
        "commit_id": merged_commit.id
    }

from server.utils.diff_utils import apply_diff

@router.post("/{repo_id}/merge/{commit_id}")
def merge_commit(
    repo_id: int,
    commit_id: int,
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    #   Access check
    access = db.query(AccessControl).filter_by(user_id=current_user.id, repository_id=repo_id).first()
    if not access or access.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can merge commits.")

    #   Fetch commit
    commit = db.query(Commit).filter_by(id=commit_id, repo_id=repo_id).first()
    if not commit or commit.status != "proposed":
        raise HTTPException(status_code=404, detail="Commit not found or already merged")

    #   Get original path
    base_filename = commit.original_filename
    if not base_filename:
        raise HTTPException(status_code=400, detail="Original filename missing")

    version_dir = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}/versions"
    os.makedirs(version_dir, exist_ok=True)
    repo_root = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}"
    #   Prepare merged content
    content = ""
    if commit.diff_text:
        original_path = os.path.join(repo_root, commit.original_filename)
        base_content = ""
        if os.path.exists(original_path):
            with open(original_path, "r", encoding="utf-8") as f:
                base_content = f.read()
        content = apply_diff(base_content, commit.diff_text)

    elif commit.snapshot_path and os.path.exists(commit.snapshot_path):
        with open(commit.snapshot_path, "r", encoding="utf-8") as f:
            content = f.read()
    else:
        raise HTTPException(status_code=400, detail="No valid snapshot or diff to merge.")

    #   Save new version
    versioned_filename = get_next_flat_version_filename(version_dir, base_filename)
    version_path = os.path.join(version_dir, versioned_filename)

    with open(version_path, "w", encoding="utf-8") as f:
        f.write(content)


    #   Mark commit as merged
    commit.status = "merged"
    commit.versioned_filename = versioned_filename
    db.commit()

    #   Log
    log = Log(
        user_id=current_user.id,
        repo_id=repo_id,
        commit_id=commit.id,
        action="merge",
        description=f"Merged commit {commit.id} into {versioned_filename}",
        timestamp=datetime.utcnow()
    )
    db.add(log)
    db.commit()
    # 1. Determine or create FileVersion row
    fv = db.query(FileVersion).filter_by(repo_id=repo_id, original_filename=base_filename).first()
    if not fv:
        fv = FileVersion(repo_id=repo_id, original_filename=base_filename, latest_version=1)
        db.add(fv)
        db.commit()
        db.refresh(fv)
    else:
        fv.latest_version += 1
        db.commit()

    # 2. Add to FileVersionHistory
    fvh = FileVersionHistory(
        repo_id=repo_id,
        original_filename=base_filename,
        version_no=fv.latest_version,
        commit_id=commit.id,
        versioned_filename=versioned_filename,
        snapshot_path=version_path,
        merged_by=current_user.id,
    )
    db.add(fvh)
    db.commit()

    return {"message": f"Commit merged as {versioned_filename}", "version_file": versioned_filename}
from server.utils.diff_utils import apply_diff

@router.get("/{repo_id}/commit_preview/{commit_id}")
def preview_commit(
    repo_id: int,
    commit_id: int,
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    commit = db.query(Commit).filter_by(id=commit_id, repo_id=repo_id).first()
    if not commit:
        raise HTTPException(status_code=404, detail="Commit not found")

    base_path = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}"
    full_path = os.path.join(base_path, commit.original_filename)

    # 🧠 1. Reconstruct content
    if commit.diff_text:
        base_content = ""
        if os.path.exists(full_path):
            with open(full_path, "r", encoding="utf-8") as f:
                base_content = f.read()

        try:
            content = apply_diff(base_content, commit.diff_text)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Diff application failed: {e}")
    elif commit.snapshot_path and os.path.exists(commit.snapshot_path):
        with open(commit.snapshot_path, "r", encoding="utf-8") as f:
            content = f.read()
    else:
        raise HTTPException(status_code=404, detail="Nothing to preview")

    return {"content": content, "filename": commit.original_filename}
@router.get("/{repo_id}/file_merge_info")
def file_merge_info(
    repo_id: int,
    file_name: str,
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    access = db.query(AccessControl).filter_by(user_id=current_user.id, repository_id=repo_id).first()
    if not access:
        raise HTTPException(status_code=403, detail="Permission denied")

    fvh = db.query(FileVersionHistory).filter_by(
        repo_id=repo_id,
        versioned_filename=file_name
    ).first()

    if not fvh:
        return {"is_merged": False}

    commit = db.query(Commit).filter_by(id=fvh.commit_id, status="merged").first()
    if commit:
        return {
            "is_merged": True,
            "commit_id": fvh.commit_id
        }

    return {"is_merged": False}


@router.get("/{repo_id}/download_version")
def download_version(
    repo_id: int,
    file: str,
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)  
):
    access = db.query(AccessControl).filter_by(user_id=current_user.id, repository_id=repo_id).first()
    if not access:
        raise HTTPException(status_code=403, detail="Permission denied")
    safe_file = os.path.basename(file)
    path = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}/versions/{safe_file}"
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(path, filename=safe_file)


@router.get("/{repo_id}/edit_file_text")
def convert_to_editable_text(
    repo_id: int,
    name: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    access = db.query(AccessControl).filter_by(user_id=current_user.id, repository_id=repo_id).first()
    if not access or access.role not in ("admin", "editor", "collaborator"):
        raise HTTPException(status_code=403, detail="Permission denied")

    file_path = os.path.join(f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}", name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Original file not found")

    from server.utils.file_parser import extract_text_from_file

    # Generate .txt version
    text_version_path = os.path.splitext(file_path)[0] + ".txt"
    try:
        content = extract_text_from_file(file_path)

        # Save the extracted text to .txt file
        with open(text_version_path, "w", encoding="utf-8") as f:
            f.write(content)

        rel_txt_path = os.path.relpath(text_version_path, f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}")
        return {
            "message": "Editable text file created",
            "editable_path": rel_txt_path,
            "content": content
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to extract and save text: {str(e)}")

@router.get("/{repo_id}/convert_version")
def convert_versioned_file(
    repo_id: int,
    filename: str,
    target_format: str,
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    #   1. Access check
    access = db.query(AccessControl).filter_by(
        user_id=current_user.id, repository_id=repo_id
    ).first()
    if not access:
        raise HTTPException(status_code=403, detail="Permission denied")

    #   2. Ensure the file exists
    repo_folder = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}"
    file_path = os.path.join(repo_folder, filename)

    if not os.path.isfile(file_path):
        raise HTTPException(status_code=404, detail="Source .txt file not found")

    #   3. Must be a .txt file
    if not filename.lower().endswith(".txt"):
        raise HTTPException(status_code=400, detail="Only .txt files can be converted")

    #   4. Read file content
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read file: {str(e)}")

    #   5. Build target path
    base_name, _ = os.path.splitext(filename)
    target_format = target_format.lower()
    if target_format not in ["docx", "pdf","doc"]:
        raise HTTPException(status_code=400, detail="Unsupported target format")

    out_path = os.path.join(repo_folder, f"{base_name}.{target_format}")

    #   6. Convert
    try:
        if target_format == "docx":
            doc = Document()
            for line in content.splitlines():
                doc.add_paragraph(line)
            doc.save(out_path)

        elif target_format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            for line in content.splitlines():
                pdf.multi_cell(0, 10, txt=line)
            pdf.output(out_path)
        elif target_format == "doc":
            doc = Document()
            for line in content.splitlines():
                doc.add_paragraph(line)
            
            # Save as .doc by renaming .docx file (Word supports it)
            temp_path = os.path.join(repo_folder, f"{base_name}.docx")
            doc.save(temp_path)
            os.rename(temp_path, out_path)  # Rename to .doc

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

    #   7. Return the converted file
    media_type = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "doc": "application/msword"
    }.get(target_format, "application/octet-stream")

    return FileResponse(out_path, media_type=media_type, filename=os.path.basename(out_path))

@router.get("/{repo_id}/version_files")
def list_version_files(
    repo_id: int,
    current_user: User = Depends(get_current_user)
):
    version_dir = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}/versions"
    if not os.path.exists(version_dir):
        return []

    return [f for f in os.listdir(version_dir) if f.endswith(".txt")]
@router.post("/{repo_id}/merge_versions")
def merge_versions(
    repo_id: int,
    payload: dict,
    db: Session = Depends(get_db),
    current_user: UserOut = Depends(get_current_user)
):
    from fastapi.responses import FileResponse
    from server.utils.diff_utils import generate_diff, apply_diff

    version_files = payload.get("version_files")
    if not version_files or len(version_files) != 2:
        raise HTTPException(status_code=400, detail="Exactly 2 version files must be selected")

    # ✅ Access check
    access = db.query(AccessControl).filter_by(
        user_id=current_user.id,
        repository_id=repo_id
    ).first()
    if not access or access.role not in ("admin", "editor","collaborator"):
        raise HTTPException(status_code=403, detail="Permission denied")

    version_dir = f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}/versions"
    os.makedirs(version_dir, exist_ok=True)

    # ✅ Resolve paths
    paths = [os.path.join(version_dir, f) for f in version_files]
    for path in paths:
        if not os.path.exists(path):
            raise HTTPException(status_code=404, detail=f"Version file '{path}' not found")

    # ✅ Read contents
    contents = []
    for path in paths:
        with open(path, "r", encoding="utf-8") as f:
            contents.append(f.read())

    # ✅ Base comparison logic
    base1 = strip_version_suffix(version_files[0])
    base2 = strip_version_suffix(version_files[1])

    if base1 == base2:
        base_filename = base1
        original_path = os.path.join(f"D:/VCS_Storage/user_{current_user.id}/repo_{repo_id}", base_filename)

        original_content = ""
        if os.path.exists(original_path):
            with open(original_path, "r", encoding="utf-8") as f:
                original_content = f.read()

        # ✅ Apply diffs to original content sequentially
        merged_versions = []
        for content in contents:
            diff = generate_diff(original_content, content)
            merged = apply_diff(original_content, diff)
            merged_versions.append(merged)

        # ✅ Merge diffs line-by-line without duplication
        lines_set = set()
        merged_lines = []
        for text in merged_versions:
            for line in text.splitlines():
                clean_line = line.strip()
                if clean_line and clean_line not in lines_set:
                    lines_set.add(clean_line)
                    merged_lines.append(clean_line)

        merged_content = "\n".join(merged_lines)

    else:
        # ✅ If base file differs — combine both without duplication
        lines_set = set()
        merged_lines = []
        for content in contents:
            for line in content.splitlines():
                clean_line = line.strip()
                if clean_line and clean_line not in lines_set:
                    lines_set.add(clean_line)
                    merged_lines.append(clean_line)

        merged_content = "\n".join(merged_lines)
        base_filename = "merged_from_" + os.path.splitext(version_files[0])[0]+".txt"

    if not merged_content.strip():
        raise HTTPException(status_code=400, detail="Merged content is empty.")

    # ✅ Save to version directory
    merged_filename = get_next_flat_version_filename(version_dir, base_filename)
    merged_path = os.path.join(version_dir, merged_filename)
    with open(merged_path, "w", encoding="utf-8") as f:
        f.write(merged_content)

    # ✅ Log commit
    merged_commit = Commit(
        repo_id=repo_id,
        author_id=current_user.id,
        message=f"Merged versions {version_files[0]} and {version_files[1]}",
        original_filename=base_filename,
        versioned_filename=merged_filename,
        snapshot_path=merged_path,
        status="merged"
    )
    db.add(merged_commit)
    db.commit()
    db.refresh(merged_commit)

    db.add(Log(
        user_id=current_user.id,
        repo_id=repo_id,
        commit_id=merged_commit.id,
        action="merge",
        description=f"Versions merged: {version_files}",
        timestamp=datetime.utcnow()
    ))
    db.commit()
    # 1. Determine or create FileVersion row
    fv = db.query(FileVersion).filter_by(repo_id=repo_id, original_filename=base_filename).first()
    if not fv:
        fv = FileVersion(repo_id=repo_id, original_filename=base_filename, latest_version=1)
        db.add(fv)
        db.commit()
        db.refresh(fv)
    else:
        fv.latest_version += 1
        db.commit()

    # 2. Add to FileVersionHistory
    fvh = FileVersionHistory(
        repo_id=repo_id,
        original_filename=base_filename,
        version_no=fv.latest_version,
        commit_id=merged_commit.id,
        versioned_filename=merged_filename,
        snapshot_path=merged_path,
        merged_by=current_user.id,
    )
    db.add(fvh)
    db.commit()
    # Optional: Get original commit ids if you track this
    commit1 = db.query(Commit).filter_by(versioned_filename=version_files[0]).first()
    commit2 = db.query(Commit).filter_by(versioned_filename=version_files[1]).first()

    if commit1 and commit2:
        merge_entry = MergeHistory(
            repo_id=repo_id,
            base_commit_id=commit1.id,
            merged_commit_id=commit2.id,
            result_commit_id=merged_commit.id,
            result_filename=merged_filename,
            merged_by=current_user.id,
            timestamp=datetime.utcnow()
        )
        db.add(merge_entry)
        db.commit()

    return {
        "message": f"Versions merged into {merged_filename}",
        "versioned_filename": merged_filename
    }
